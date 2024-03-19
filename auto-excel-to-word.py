import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
import tkinter as tk
from tkinter import filedialog
import re
from tkinter import *
import tkinter.messagebox
 
def getInput(title, message):
  def return_callback(event):
    print('quit...')
    root.quit()
  def close_callback():
    tkinter.messageBox.showinfo('message', 'no click...')
  root = Tk(className=title)
  root.wm_attributes('-topmost', 1)
  screenwidth, screenheight = root.maxsize()
  width = 300
  height = 100
  size = '%dx%d+%d+%d' % (width, height, (screenwidth - width)/2, (screenheight - height)/2)
  root.geometry(size)
  root.resizable(0, 0)
  lable = Label(root, height=2)
  lable['text'] = message
  lable.pack()
  entry = Entry(root)
  entry.bind('<Return>', return_callback)
  entry.pack()
  entry.focus_set()
  root.protocol("WM_DELETE_WINDOW", close_callback)
  root.mainloop()
  str = entry.get()
  root.destroy()
  return str

# *一些基础设置与函数
def number(str_list):
    # *将一开始的编号删除
    test=list()
    for i in range(len(str_list)):
        for j in range(len(str_list[i])):
            pattern=re.compile(r'^[0-9].')
            test.append(pattern.sub("",str_list[i][j]))
    return test
def addnumber(listresult,doc,title,paragraphnum,trueorfalse):
    # *将对应数值按照列开始放入
    result=list()
    paragraph1 = doc.add_paragraph()  # 创建一个段落对象
    run1 = paragraph1.add_run(title)
    run1.font.bold = True
    for i in range(len(listresult)):
        if trueorfalse:
            doc.add_paragraph(str(paragraphnum)+'.'+str(i+1)+'.'+listresult[i])
        else:
            doc.add_paragraph(str(i+1)+'.'+listresult[i])
def readdata(dict_list,number):
    # * 最开始读取数据的位置，将某一列的数值进行读取，然后将每一条目分开变成数组,分开的数组为missions
    mission=list()
    missionwithname=list()
    missions=list()
    for i in range(len(people)):
        if dict_list[2+i]['Unnamed: '+str(number)]!=None:
            mission.append(dict_list[2+i]['Unnamed: '+str(number)])
        #这里放number3与5
    print(mission)
    for i in range(len(mission)):
        missionwithname.append(re.sub('\n','【'+people[i]+'】'+'\n',mission[i])+'【'+people[i]+'】')
    for i in range(len(missionwithname)):
        missions.append(re.split('\n',missionwithname[i]))
        # for j in range(len(mission[i])):
        #     missionwithname.append(mission[i])
        # mission[i].sub('\n',people[i])
    return missions
    # *这里就变成了split
def splitbyprandfengzhuang(missions): 
    # *将pr和封装分开放进去
    fengzhuang=list()
    pr=list()
    for i in range(len(missions)):
        if i==0 or i==3 or i==7:
            fengzhuang.append(missions[i]) 
        else:
            pr.append(missions[i])
    return fengzhuang,pr

if __name__ == '__main__':
    doc=Document()
    doc.styles['Normal'].font.name = u'仿宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    doc.styles['Normal'].font.size = Pt(11.5)
    doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    root = tk.Tk()
    root.withdraw()
    # 获取文件夹路径
    f_path = filedialog.askopenfilename()
    Folderpath = filedialog.askdirectory() #获得选择好的文件夹
    print(Folderpath)
    # * 这里开始使用
    data = pd.read_excel(f_path)
    dict_list = data.to_dict(orient="records")
    txt=getInput('auto-excel-to-word','输入你想给文件的命名')
    """
    data_dict[2]开始是people1
    """
    wordname=Folderpath+'/'+txt+'.docx'
    people=['people1','people2','people3','people4','people5','people6','people7','people8']
    # 这里添加的上周工作小结
    missions=readdata(dict_list,3)
    fengzhuang,pr=splitbyprandfengzhuang(missions)
    listpr=number(pr)
    listfengzhuang=number(fengzhuang)
    doc.add_paragraph("本周完成工作")
    doc.add_paragraph("1、xx方向")
    addnumber(listpr,doc,'xx与xxxx脚本调试方面：',1,0)
    addnumber(listfengzhuang,doc,'封x工作方面：',1,0)
    doc.save(wordname)
    missions=readdata(dict_list,5)
    fengzhuang,pr=splitbyprandfengzhuang(missions)
    listpr=number(pr)
    listfengzhuang=number(fengzhuang)
    # 注意不能有空白段落
    doc.add_paragraph("下周计划")
    doc.add_paragraph("1、xx方向")
    addnumber(listpr,doc,'xx与xxxx脚本调试方面：',1,0)
    addnumber(listfengzhuang,doc,'封x工作方面：',1,0)
    doc.save(wordname)